"""Clean-docx builder for NIH attachments.

Produces a fresh .docx with consistent formatting: Arial 11, 0.5-inch margins,
justified paragraphs, compact spacing — the standard 1-page NIH response format
with bold section labels, italic paraphrases, and normal body text.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from docx import Document
from docx.oxml.ns import qn as docx_qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Alignment shorthand accepted in block specs
_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# Type alias: a block is either a plain string or a dict spec
Block = Union[str, Dict[str, Any]]


def heading(text: str) -> Dict[str, Any]:
    """Return a centered bold heading block dict."""
    return {
        "runs": [{"text": text, "bold": True}],
        "align": "center",
    }


def body(*runs: Union[str, Dict[str, Any]]) -> Dict[str, Any]:
    """Return a justified body block dict from run specs or plain strings.

    Each positional arg can be:
    - a str  => plain normal run
    - a dict => run spec with keys ``text``, ``bold``, ``italic``
    """
    run_list: List[Dict[str, Any]] = []
    for r in runs:
        if isinstance(r, str):
            run_list.append({"text": r})
        else:
            run_list.append(r)
    return {"runs": run_list}


def new_doc(
    out_path: Union[str, Path],
    blocks: List[Block],
    *,
    font: str = "Arial",
    size_pt: int = 11,
    margin_in: float = 0.5,
    justify: bool = True,
    compact: bool = True,
) -> Path:
    """Build a new .docx from *blocks* and write it to *out_path*.

    Parameters
    ----------
    out_path:
        Destination path for the .docx file.
    blocks:
        Ordered list of paragraph specs.  Each element is either:

        * A **plain string** — rendered as a single normal run with default
          alignment.
        * A **dict** with keys:

          ``runs``
              List of run dicts, each with ``text`` (required), ``bold``
              (bool, default False), ``italic`` (bool, default False).
          ``align``
              ``"left"``, ``"center"``, or ``"both"`` / ``"justify"``.
              Defaults to JUSTIFY when *justify* is True.
          ``space_after_pt``
              Override the compact space-after (points) for this block.

    font:
        Font family applied to the Normal style and every run.
    size_pt:
        Font size in points.
    margin_in:
        Page margin (inches) applied to all four sides of every section.
    justify:
        When True, default paragraph alignment is JUSTIFY.
    compact:
        When True, default space-after is 4 pt (unless *space_after_pt*
        overrides it per block).

    Returns
    -------
    Path
        The written file path.
    """
    out_path = Path(out_path)
    doc = Document()

    # --- Normal style: font + size ----------------------------------------
    normal = doc.styles["Normal"]
    normal_font = normal.font
    normal_font.name = font
    normal_font.size = Pt(size_pt)

    # Also patch the underlying rPr so asian/complex fonts don't override.
    rpr = normal.element.get_or_add_rPr()
    for tag in (docx_qn("w:rFonts"),):
        existing = rpr.find(tag)
        if existing is not None:
            rpr.remove(existing)
    from lxml import etree
    r_fonts = etree.SubElement(rpr, docx_qn("w:rFonts"))
    r_fonts.set(docx_qn("w:ascii"), font)
    r_fonts.set(docx_qn("w:hAnsi"), font)
    r_fonts.set(docx_qn("w:cs"), font)

    # --- Sections: margins ------------------------------------------------
    margin = Inches(margin_in)
    for section in doc.sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

    # --- Default paragraph format -----------------------------------------
    default_align = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    default_space_after = Pt(4) if compact else Pt(8)

    # --- Paragraphs -------------------------------------------------------
    for block in blocks:
        # Normalise to dict
        if isinstance(block, str):
            spec: Dict[str, Any] = {"runs": [{"text": block}]}
        else:
            spec = block

        run_specs: List[Dict[str, Any]] = spec.get("runs", [])
        align_key: Optional[str] = spec.get("align")
        space_after_pt: Optional[int] = spec.get("space_after_pt")

        para = doc.add_paragraph()

        # Alignment
        if align_key:
            para.alignment = _ALIGN.get(align_key, default_align)
        else:
            para.alignment = default_align

        # Paragraph format: space after
        pf = para.paragraph_format
        if space_after_pt is not None:
            pf.space_after = Pt(space_after_pt)
        else:
            pf.space_after = default_space_after
        pf.space_before = Pt(0)
        pf.line_spacing = None  # keep single / auto

        # Runs
        for rs in run_specs:
            run = para.add_run(rs.get("text", ""))
            run.bold = rs.get("bold", False)
            run.italic = rs.get("italic", False)
            run.font.name = font
            run.font.size = Pt(size_pt)

    doc.save(str(out_path))
    return out_path
