"""Tests for zoterocite.builder — new-doc factory for NIH attachments."""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches

from zoterocite.builder import new_doc, heading, body
from zoterocite.docxio import Docx
from zoterocite.views import read_views, counts
from zoterocite.ooxml import qn


@pytest.fixture()
def built_doc(tmp_path: Path) -> Path:
    """Build a representative test document and return its path."""
    out = tmp_path / "nih_test.docx"
    blocks = [
        heading("Specific Aims Response"),
        body(
            {"text": "Reviewer concern: ", "bold": True},
            {"text": "the approach is unvalidated.", "italic": True},
            {"text": " We have addressed this fully."},
        ),
        body(
            {"text": "Our preliminary data demonstrate robustness across cohorts."},
        ),
    ]
    new_doc(out, blocks)
    return out


# ---------------------------------------------------------------------------
# Text content checks
# ---------------------------------------------------------------------------

def test_text_content(built_doc: Path) -> None:
    """All run texts must appear in the accepted view."""
    views = read_views(built_doc)
    accepted = views["accepted"]
    assert "Specific Aims Response" in accepted
    assert "Reviewer concern:" in accepted
    assert "the approach is unvalidated." in accepted
    assert "We have addressed this fully." in accepted
    assert "Our preliminary data demonstrate robustness across cohorts." in accepted


def test_word_count_positive(built_doc: Path) -> None:
    """counts() must report at least one word."""
    views = read_views(built_doc)
    result = counts(views["accepted"])
    assert result["words"] > 0


# ---------------------------------------------------------------------------
# Margin check (via python-docx and via raw OOXML)
# ---------------------------------------------------------------------------

_TWIPS_PER_INCH = 1440  # 1 inch = 1440 twips (EMU-free unit Word uses for pgMar)
_HALF_INCH_TWIPS = int(0.5 * _TWIPS_PER_INCH)   # 720
_TOLERANCE = 5  # allow ±5 twips rounding


def test_section_margin_python_docx(built_doc: Path) -> None:
    """python-docx reports 0.5-inch margins on the first section."""
    doc = Document(str(built_doc))
    sec = doc.sections[0]
    assert sec.left_margin == Inches(0.5)
    assert sec.right_margin == Inches(0.5)
    assert sec.top_margin == Inches(0.5)
    assert sec.bottom_margin == Inches(0.5)


def test_section_margin_ooxml(built_doc: Path) -> None:
    """sectPr/pgMar w:left attribute is ~720 twips (0.5 in)."""
    root = Docx(built_doc).read_tree()
    # sectPr can appear inside the body or as the last child of body
    body_el = root.find(".//" + qn("w:body"))
    assert body_el is not None, "No w:body found"
    sect_pr = body_el.find(qn("w:sectPr"))
    assert sect_pr is not None, "No w:sectPr found"
    pg_mar = sect_pr.find(qn("w:pgMar"))
    assert pg_mar is not None, "No w:pgMar found"
    left_val = int(pg_mar.get(qn("w:left"), "0"))
    assert abs(left_val - _HALF_INCH_TWIPS) <= _TOLERANCE, (
        f"Expected ~{_HALF_INCH_TWIPS} twips, got {left_val}"
    )


# ---------------------------------------------------------------------------
# Font checks
# ---------------------------------------------------------------------------

def test_document_font_arial(built_doc: Path) -> None:
    """At least one run in the document has font name Arial."""
    doc = Document(str(built_doc))
    fonts_seen = {
        run.font.name
        for para in doc.paragraphs
        for run in para.runs
        if run.font.name
    }
    assert "Arial" in fonts_seen, f"Arial not found; fonts seen: {fonts_seen}"


# ---------------------------------------------------------------------------
# Formatting checks
# ---------------------------------------------------------------------------

def test_run_bold_italic(built_doc: Path) -> None:
    """Bold and italic attributes are preserved on the relevant runs."""
    doc = Document(str(built_doc))
    paras = [p for p in doc.paragraphs if p.runs]
    # second paragraph (index 1) has bold + italic + normal runs
    mixed_para = paras[1]
    bold_runs = [r for r in mixed_para.runs if r.bold]
    italic_runs = [r for r in mixed_para.runs if r.italic]
    assert bold_runs, "Expected a bold run in the second paragraph"
    assert italic_runs, "Expected an italic run in the second paragraph"


def test_heading_centered(built_doc: Path) -> None:
    """The heading paragraph is center-aligned."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document(str(built_doc))
    first_para = doc.paragraphs[0]
    assert first_para.alignment == WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Convenience helpers
# ---------------------------------------------------------------------------

def test_heading_helper() -> None:
    spec = heading("Title")
    assert spec["runs"][0]["bold"] is True
    assert spec["runs"][0]["text"] == "Title"
    assert spec["align"] == "center"


def test_body_helper_strings() -> None:
    spec = body("Hello", "World")
    assert len(spec["runs"]) == 2
    assert spec["runs"][0] == {"text": "Hello"}


def test_body_helper_dicts() -> None:
    spec = body({"text": "Bold", "bold": True}, "normal")
    assert spec["runs"][0]["bold"] is True
    assert spec["runs"][1] == {"text": "normal"}


def test_string_block(tmp_path: Path) -> None:
    """Plain-string blocks are accepted and produce text in the document."""
    out = tmp_path / "plain.docx"
    new_doc(out, ["Just a plain sentence."])
    views = read_views(out)
    assert "Just a plain sentence." in views["accepted"]


def test_space_after_override(tmp_path: Path) -> None:
    """space_after_pt block override is applied."""
    out = tmp_path / "space.docx"
    new_doc(out, [{"runs": [{"text": "Big gap"}], "space_after_pt": 24}])
    doc = Document(str(out))
    from docx.shared import Pt
    para = doc.paragraphs[0]
    assert para.paragraph_format.space_after == Pt(24)
