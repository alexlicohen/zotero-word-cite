"""Tests for zoterocite.paras.paragraph_text — the run-reassembly primitive.

Pins the accepted-view text model: tracked deletions excluded, and
flow-affecting children (``<w:br>``/``<w:cr>``/``<w:tab>``) rendered as
separators rather than silently dropped (which used to glue adjacent words).
"""
import pytest
from lxml import etree

from zoterocite.ooxml import qn
from zoterocite.paras import (
    ParaIndex,
    _norm,
    find_paragraph,
    find_paragraphs,
    iter_paragraphs,
    paragraph_text,
)
from zoterocite.views import read_views


def _run(parent, text):
    r = etree.SubElement(parent, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.text = text
    return r


def _break(parent, kind=None):
    r = etree.SubElement(parent, qn("w:r"))
    b = etree.SubElement(r, qn("w:br"))
    if kind:
        b.set(qn("w:type"), kind)
    return r


def _tab(parent):
    r = etree.SubElement(parent, qn("w:r"))
    etree.SubElement(r, qn("w:tab"))
    return r


def test_plain_runs_concatenate():
    # Adjacent <w:t> runs with no break between them stay glued (a word split
    # across runs must NOT gain a spurious separator).
    p = etree.Element(qn("w:p"))
    _run(p, "neuro")
    _run(p, "logy")
    assert paragraph_text(p) == "neurology"


def test_soft_break_renders_as_newline():
    p = etree.Element(qn("w:p"))
    _run(p, "Line one")
    _break(p)
    _run(p, "Line two")
    assert paragraph_text(p) == "Line one\nLine two"


def test_tab_and_page_break_render():
    p = etree.Element(qn("w:p"))
    _run(p, "a")
    _tab(p)
    _run(p, "b")
    _break(p, "page")
    _run(p, "c")
    assert paragraph_text(p) == "a\tb\nc"


def test_carriage_return_renders_as_newline():
    p = etree.Element(qn("w:p"))
    _run(p, "x")
    r = etree.SubElement(p, qn("w:r"))
    etree.SubElement(r, qn("w:cr"))
    _run(p, "y")
    assert paragraph_text(p) == "x\ny"


def test_tracked_deletion_excluded():
    p = etree.Element(qn("w:p"))
    _run(p, "keep ")
    d = etree.SubElement(p, qn("w:del"))
    dr = etree.SubElement(d, qn("w:r"))
    dt = etree.SubElement(dr, qn("w:delText"))
    dt.text = "GONE"
    _run(p, " tail")
    assert paragraph_text(p) == "keep  tail"


def test_literal_newline_in_wt_preserved():
    # python-docx add_run("a\nb") stores a literal newline inside one <w:t>;
    # it must survive unchanged (not be doubled or stripped).
    p = etree.Element(qn("w:p"))
    r = etree.SubElement(p, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.text = "1. a\n2. b"
    assert paragraph_text(p) == "1. a\n2. b"


def test_matches_read_views_accepted_content(tmp_path):
    # A single-paragraph doc whose only paragraph carries a soft break: the
    # per-w:p paragraph_text content equals the read_views accepted content.
    from docx import Document

    doc = Document()
    para = doc.add_paragraph()
    para.add_run("first")
    para.add_run().add_break()  # real <w:br/>
    para.add_run("second")
    out = tmp_path / "br.docx"
    doc.save(str(out))

    from zoterocite.docxio import Docx, DOCUMENT
    from zoterocite.paras import iter_paragraphs

    root = Docx(str(out)).read_tree(DOCUMENT)
    body_paras = [p for p in iter_paragraphs(root) if paragraph_text(p)]
    assert body_paras and paragraph_text(body_paras[0]) == "first\nsecond"
    assert read_views(str(out))["accepted"] == "first\nsecond"


# ---------------------------------------------------------------------------
# ParaIndex equivalence — the cached one-shot resolver must be BYTE-IDENTICAL to
# the old per-call find_paragraph(s). A reference implementation below pins the
# OLD semantics independently of the (now-delegating) module functions, so the
# equivalence assertion has teeth even though find_paragraph(s) call ParaIndex.
# ---------------------------------------------------------------------------

def _ref_find_all(root, anchor):
    """The pre-ParaIndex find_paragraphs, verbatim (the contract under test)."""
    a = _norm(anchor)
    return [p for p in iter_paragraphs(root) if a in _norm(paragraph_text(p))]


def _ref_find(root, anchor):
    hits = _ref_find_all(root, anchor)
    if len(hits) != 1:
        raise LookupError(
            f"anchor {anchor!r} matched {len(hits)} paragraphs (need exactly 1)")
    return hits[0]


def _para(body, *runs):
    """Append a <w:p> whose visible text is split across several <w:r><w:t> runs."""
    p = etree.SubElement(body, qn("w:p"))
    for text in runs:
        r = etree.SubElement(p, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = text
    return p


def _textbox_para(body, caption):
    """A body <w:p> holding an mc:AlternateContent textbox: the caption text lives
    in the body paragraph (read via the mc:Choice) AND is duplicated in nested
    Choice/Fallback w:txbxContent paragraphs that iter_paragraphs must exclude."""
    p = etree.SubElement(body, qn("w:p"))
    alt = etree.SubElement(p, qn("mc:AlternateContent"))
    for branch in ("mc:Choice", "mc:Fallback"):
        b = etree.SubElement(alt, qn(branch))
        txbx = etree.SubElement(b, qn("w:txbxContent"))
        inner_p = etree.SubElement(txbx, qn("w:p"))
        r = etree.SubElement(inner_p, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = caption
    return p


def _equivalence_root():
    root = etree.Element(qn("w:document"))
    body = etree.SubElement(root, qn("w:body"))
    paras = {}
    # run-split sentence
    paras["split"] = _para(body, "lesion ", "network ", "mapping ", "is robust")
    # whitespace + tab variant (paragraph_text renders <w:tab> as "\t"; _norm collapses)
    tab_p = etree.SubElement(body, qn("w:p"))
    r1 = etree.SubElement(tab_p, qn("w:r")); etree.SubElement(r1, qn("w:t")).text = "hello"
    rt = etree.SubElement(tab_p, qn("w:r")); etree.SubElement(rt, qn("w:tab"))
    r2 = etree.SubElement(tab_p, qn("w:r")); etree.SubElement(r2, qn("w:t")).text = "world"
    paras["tab"] = tab_p
    # ambiguous: two paragraphs sharing a substring, document order preserved
    paras["dupA"] = _para(body, "shared marker (12)")
    paras["dupB"] = _para(body, "again shared marker (12) here")
    # tracked deletion excluded from the accepted-view text
    del_p = etree.SubElement(body, qn("w:p"))
    rk = etree.SubElement(del_p, qn("w:r")); etree.SubElement(rk, qn("w:t")).text = "keep "
    d = etree.SubElement(del_p, qn("w:del"))
    dr = etree.SubElement(d, qn("w:r")); etree.SubElement(dr, qn("w:delText")).text = "STRUCK"
    rt2 = etree.SubElement(del_p, qn("w:r")); etree.SubElement(rt2, qn("w:t")).text = " tail"
    paras["del"] = del_p
    # textbox: caption duplicated in nested txbxContent paragraphs (must NOT be yielded)
    paras["textbox"] = _textbox_para(body, "Figure caption alpha")
    return root, paras


ANCHORS = [
    "lesion network mapping",        # unique, spans run boundaries
    "hello world",                   # unique, tab collapsed by _norm
    "hello\tworld",                  # same, anchor itself carries the tab
    "shared marker (12)",            # AMBIGUOUS (two paragraphs, doc order)
    "STRUCK",                        # NOT FOUND (lives under w:del)
    "keep  tail",                    # unique, deletion gap preserved
    "Figure caption alpha",          # textbox caption -> body p only, not the 2 inner
    "no such anchor anywhere",       # NOT FOUND
    "  lesion   network  mapping  ", # whitespace-normalized to the unique anchor
]


@pytest.mark.parametrize("anchor", ANCHORS)
def test_paraindex_find_all_matches_reference(anchor):
    root, _ = _equivalence_root()
    idx = ParaIndex(root)
    ref = _ref_find_all(root, anchor)
    got = idx.find_all(anchor)
    # Same elements, same first-appearance ORDER (identity, not just text equality).
    assert [id(p) for p in got] == [id(p) for p in ref]
    # And the public delegator agrees with the index it now wraps.
    assert [id(p) for p in find_paragraphs(root, anchor)] == [id(p) for p in ref]


@pytest.mark.parametrize("anchor", ANCHORS)
def test_paraindex_find_matches_reference(anchor):
    root, _ = _equivalence_root()
    idx = ParaIndex(root)
    ref_hits = _ref_find_all(root, anchor)
    if len(ref_hits) == 1:
        assert idx.find(anchor) is ref_hits[0]
        assert find_paragraph(root, anchor) is ref_hits[0]
    else:
        # ambiguous (>1) and not-found (0) both raise the IDENTICAL LookupError.
        with pytest.raises(LookupError) as ei_ref:
            _ref_find(root, anchor)
        with pytest.raises(LookupError) as ei_idx:
            idx.find(anchor)
        with pytest.raises(LookupError) as ei_pub:
            find_paragraph(root, anchor)
        assert str(ei_idx.value) == str(ei_ref.value) == str(ei_pub.value)


def test_paraindex_excludes_textbox_inner_paragraphs():
    # The caption resolves to exactly the ONE body paragraph; the duplicated
    # Choice/Fallback w:txbxContent paragraphs are never returned (no "matched 3").
    root, paras = _equivalence_root()
    idx = ParaIndex(root)
    hits = idx.find_all("Figure caption alpha")
    assert hits == [paras["textbox"]]
    assert idx.find("Figure caption alpha") is paras["textbox"]


def test_tracked_move_source_excluded():
    # A tracked-move SOURCE (<w:moveFrom>) uses ordinary <w:t>, so the moved text
    # would otherwise appear at BOTH source and destination. The accepted view drops
    # the source (it moved away); paragraph_text must mirror views._extract and
    # exclude it. A moveTo DESTINATION is kept (survives in the accepted view).
    p = etree.Element(qn("w:p"))
    _run(p, "before ")
    mf = etree.SubElement(p, qn("w:moveFrom"))
    mf.set(qn("w:id"), "1")
    mf.set(qn("w:author"), "Collaborator, A.")
    mf.set(qn("w:date"), "2026-06-20T10:00:00Z")
    mr = etree.SubElement(mf, qn("w:r"))
    etree.SubElement(mr, qn("w:t")).text = "MOVEDSOURCE "
    _run(p, "after")
    assert paragraph_text(p) == "before after"


def test_tracked_move_destination_kept():
    # A moveTo DESTINATION is part of the accepted view -> kept by paragraph_text.
    p = etree.Element(qn("w:p"))
    _run(p, "here is ")
    mt = etree.SubElement(p, qn("w:moveTo"))
    mt.set(qn("w:id"), "2")
    mr = etree.SubElement(mt, qn("w:r"))
    etree.SubElement(mr, qn("w:t")).text = "the moved text"
    assert paragraph_text(p) == "here is the moved text"


def test_move_source_excluded_matches_read_views(tmp_path):
    # End-to-end: paragraph_text of a moveFrom-bearing paragraph equals the
    # read_views accepted content (the single-owner accepted-view rendering).
    from docx import Document

    doc = Document()
    para = doc.add_paragraph()
    para.add_run("before ")
    p_el = para._p
    mf = etree.SubElement(p_el, qn("w:moveFrom"))
    mf.set(qn("w:id"), "1")
    mf.set(qn("w:author"), "Collaborator, A.")
    mf.set(qn("w:date"), "2026-06-20T10:00:00Z")
    mr = etree.SubElement(mf, qn("w:r"))
    etree.SubElement(mr, qn("w:t")).text = "MOVED "
    para.add_run("after")
    out = tmp_path / "mv.docx"
    doc.save(str(out))

    from zoterocite.docxio import Docx, DOCUMENT

    root = Docx(str(out)).read_tree(DOCUMENT)
    body_paras = [p for p in iter_paragraphs(root) if paragraph_text(p)]
    assert body_paras and paragraph_text(body_paras[0]) == "before after"
    assert read_views(str(out))["accepted"] == "before after"
